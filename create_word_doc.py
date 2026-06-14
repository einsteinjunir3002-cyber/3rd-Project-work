import docx
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import re

root_dir = r"c:\Users\einst\Desktop\PROJECT WORK 3RD PHASE"
md_path = r"c:\Users\einst\Desktop\PROJECT WORK 3RD PHASE\project_documentation.md"
docx_path = r"c:\Users\einst\Desktop\PROJECT WORK 3RD PHASE\SmartLearn_Project_Documentation.docx"

doc = docx.Document()

# Helper for shading cell background
def set_cell_background(cell, hex_color):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)

# Page Margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

def add_heading_styled(text, level):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(6)
    h.paragraph_format.keep_with_next = True
    run = h.runs[0]
    run.font.name = 'Arial'
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(99, 102, 241)  # Violet
        run.bold = True
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(79, 70, 229)  # Indigo
        run.bold = True
    else:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(55, 65, 81)   # Charcoal
        run.bold = True
    return h

def add_p_styled(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            cleaned_text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', part)
            run = p.add_run(cleaned_text)
        run.font.name = 'Arial'
        run.font.size = Pt(10.5)
    return p

# Parse markdown documentation first
with open(md_path, 'r', encoding='utf-8') as f:
    lines = f.readlines()

in_table = False
table_headers = []
table_rows = []

for line in lines:
    line = line.strip()
    if not line:
        continue
    
    if line.startswith('|'):
        parts = [p.strip() for p in line.split('|')[1:-1]]
        if not in_table:
            in_table = True
            table_headers = parts
        elif line.startswith('| :---') or line.startswith('|:---') or line.startswith('|---') or line.startswith('| ---'):
            continue
        else:
            table_rows.append(parts)
        continue
    else:
        if in_table:
            if table_headers and table_rows:
                table = doc.add_table(rows=1, cols=len(table_headers))
                table.style = 'Light Shading Accent 1'
                hdr_cells = table.rows[0].cells
                for i, h_text in enumerate(table_headers):
                    hdr_cells[i].text = h_text
                    for r in hdr_cells[i].paragraphs[0].runs:
                        r.font.name = 'Arial'
                        r.font.bold = True
                        r.font.size = Pt(10)
                for row_data in table_rows:
                    row_cells = table.add_row().cells
                    for i, val in enumerate(row_data):
                        val_cleaned = val.replace('**', '')
                        cleaned_val = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', val_cleaned)
                        row_cells[i].text = cleaned_val
                        for r in row_cells[i].paragraphs[0].runs:
                            r.font.name = 'Arial'
                            r.font.size = Pt(9.5)
            in_table = False
            table_headers = []
            table_rows = []

    if line.startswith('# '):
        add_heading_styled(line[2:], level=1)
    elif line.startswith('## '):
        add_heading_styled(line[3:], level=2)
    elif line.startswith('### '):
        add_heading_styled(line[4:], level=3)
    elif line.startswith('- ') or line.startswith('* '):
        # Determine indentation level
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        bullet_text = line[2:]
        parts = re.split(r'(\*\*.*?\*\*)', bullet_text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                cleaned_text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', part)
                run = p.add_run(cleaned_text)
            run.font.name = 'Arial'
            run.font.size = Pt(10.5)
    elif line.startswith('>') or line.startswith(':::'):
        text = line.replace('>', '').replace('[!NOTE]', '').replace('[!IMPORTANT]', '').strip()
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.italic = True
        run.font.name = 'Arial'
        run.font.size = Pt(10)
    elif line.startswith('```') or line.startswith('graph') or line.startswith('subgraph') or line.startswith('end') or line.startswith('style'):
        continue
    else:
        add_p_styled(line)

# Section 5: Append full code contents
add_heading_styled("5. Appendices: Full Code Contents", level=1)
p = doc.add_paragraph()
p.add_run("This section contains the full, unaltered code files modified during the project implementation.").font.name = 'Arial'

files_to_include = [
    ("backend/src/models/User.ts", "backend/src/models/User.ts"),
    ("backend/src/controllers/authController.ts", "backend/src/controllers/authController.ts"),
    ("backend/src/services/aiService.ts", "backend/src/services/aiService.ts"),
    ("frontend/src/types/index.ts", "frontend/src/types/index.ts"),
    ("frontend/src/pages/AuthPage.tsx", "frontend/src/pages/AuthPage.tsx"),
    ("frontend/src/components/Sidebar.tsx", "frontend/src/components/Sidebar.tsx"),
    ("frontend/src/pages/StudentHub.tsx", "frontend/src/pages/StudentHub.tsx"),
    ("frontend/src/pages/LecturerHub.tsx", "frontend/src/pages/LecturerHub.tsx"),
    ("index.html", "index.html"),
    ("assets/js/app.js", "assets/js/app.js")
]

for rel_path, title in files_to_include:
    full_path = os.path.join(root_dir, rel_path)
    add_heading_styled(f"📄 Full Code: {title}", level=2)
    
    if not os.path.exists(full_path):
        doc.add_paragraph(f"[Error: File {rel_path} not found on disk]")
        continue
        
    with open(full_path, 'r', encoding='utf-8', errors='replace') as code_f:
        code_content = code_f.read()
        
    # Put code inside a single-cell table for neat styling
    code_table = doc.add_table(rows=1, cols=1)
    code_table.autofit = False
    code_table.columns[0].width = Inches(6.5)
    cell = code_table.rows[0].cells[0]
    set_cell_background(cell, "F3F4F6")  # Light gray background
    
    code_p = cell.paragraphs[0]
    code_p.paragraph_format.space_before = Pt(4)
    code_p.paragraph_format.space_after = Pt(4)
    code_p.paragraph_format.line_spacing = 1.0
    
    run = code_p.add_run(code_content)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.0)
    run.font.color.rgb = RGBColor(31, 41, 55)  # Dark gray text

try:
    doc.save(docx_path)
    print("SUCCESS: Saved to " + docx_path)
except PermissionError:
    alt_path = docx_path.replace(".docx", "_Updated.docx")
    doc.save(alt_path)
    print("SUCCESS: Saved to alternative path due to file lock: " + alt_path)
